"""One-off generator for VibeStart technical Word doc. Run: python scripts/generate_word_doc.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_p(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_method_line(doc: Document, line: str):
    """Each line should start with # per user convention."""
    p = doc.add_paragraph(line, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def main():
    out = Path(__file__).resolve().parent.parent / "VibeStart_Technical_Documentation.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_h(doc, "VibeStart — Technical documentation", 0)
    add_p(
        doc,
        "Version 1.0 — use cases, architecture, workflow, tech stack, and code structure "
        "with one-line method descriptions prefixed with # (Markdown-style emphasis on each fact).",
    )

    # --- Use cases ---
    add_h(doc, "# Use cases", 1)
    add_p(
        doc,
        "Founders and builders capture a startup pitch as text and/or a public website URL; "
        "the system scrapes pages, stores raw HTML and excerpts locally (SQLite + files), "
        "merges context with an LLM, and optionally attaches competitor URLs for the same treatment.",
    )
    add_p(
        doc,
        "Users run AI analyses: deep competitor comparison (multi-competitor), VC discovery with match scores, "
        "hiring/job opportunities with organizational gaps, strategic gap analysis, tailored startup ideas, "
        "a skeptical 'roast', collaboration ideas, plus extra strategic modules (moats, pivots, etc.).",
    )
    add_p(
        doc,
        "All structured analysis results can be persisted as snapshots in the database for later review; "
        "in local dev, data can be cleared automatically on backend start.",
    )

    # --- Architecture ---
    add_h(doc, "# Implementation architecture", 1)
    add_p(
        doc,
        "Split stack: React 18 SPA (Vite) talks to a Node.js Express API over HTTP; "
        "development traffic is proxied from the frontend dev server to the API on 127.0.0.1.",
    )
    add_p(
        doc,
        "Persistence: better-sqlite3 (SQLite WAL) under backend/data/vibestart.sqlite; "
        "scraped HTML files under backend/data/raw_html/; Row-Level schema via foreign keys "
        "linking profiles, scraped_pages, competitors, and analysis_snapshots.",
    )
    add_p(
        doc,
        "Intelligence: OpenAI Chat Completions (JSON mode / response_format json_object) "
        "through a small service layer; scraping via axios + cheerio (server-side only).",
    )
    add_p(
        doc,
        "Security posture: API keys only on the server; CORS limited to dev front-end origins; "
        "no multi-user auth in the current MVP (single-machine / local dev assumption).",
    )
    add_p(
        doc,
        "Local-first reset behavior: by default the backend can wipe SQLite tables and raw_html cache on process start; "
        "set RESET_DB_ON_START=false in backend/.env to persist across restarts.",
    )

    # --- Workflow ---
    add_h(doc, "# Workflow (end-to-end)", 1)
    add_p(doc, "1) User submits startup text and/or URL → POST /api/profile → profile row ensured, optional scrape, LLM merge → combined_summary stored.")
    add_p(doc, "2) User adds competitor URLs → POST /api/competitors → each URL scraped, rows in scraped_pages and competitors.")
    add_p(doc, "3) User triggers an analysis → POST /api/analyze/{key} (comparison|vcs|jobs|gaps|ideas|roast|collaborations|extras) → loads profile + competitor summaries → OpenAI → snapshot row in analysis_snapshots.")
    add_p(doc, "4) Optional: GET /api/profile/:id or GET /api/snapshots/:profileId to read stored state (UI may cache profileId in localStorage).")
    add_p(doc, "5) UI hygiene: on every new run (or profile/competitor change), previous analysis panels are cleared to prevent stale mixed outputs.")

    # --- Tech flow ---
    add_h(doc, "# Tech flow (request/response + files)", 1)
    add_p(
        doc,
        "Frontend (frontend/src/App.jsx) → API helpers (frontend/src/api.js) → Vite dev proxy (frontend/vite.config.js) → "
        "Express router (backend/src/routes.js) → (optional) Scraper (backend/src/scraper.js + backend/data/raw_html/*) → "
        "OpenAI service (backend/src/openaiService.js) → JSON response → Rendered as readable markdown/cards (frontend/src/components/AnalysisResult.jsx).",
    )

    # --- Tech ---
    add_h(doc, "# Technology stack", 1)
    bullets = [
        "Frontend: React 18, Vite 6, JavaScript (ES modules), fetch API.",
        "Backend: Node.js (ESM), Express 4, dotenv, cors, better-sqlite3, axios, cheerio, openai SDK.",
        "Data: SQLite file + raw HTML artifacts on disk.",
        "Dev orchestration: concurrently + wait-on so Vite starts after /api/health is live.",
        "AI model: configurable via OPENAI_MODEL (default gpt-4o-mini).",
        "Rendering: react-markdown + remark-gfm for markdown tables and rich text output.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # --- Code structure ---
    add_h(doc, "# Repository / code structure", 1)
    add_p(
        doc,
        "Root: package.json (dev:all, wait-on), backend/ Node service, frontend/ Vite app, "
        "scripts/generate_word_doc.py (this document).",
    )

    add_h(doc, "## backend/src/index.js", 2)
    methods = [
        "# (bootstrap) — Wires CORS, JSON body parser, health route, API router, JSON parse error handler, global JSON error handler, and listens on HOST:PORT with long-lived HTTP timeouts.",
        "# GET /api/health — Returns a small JSON payload confirming the API process is running.",
    ]
    for m in methods:
        add_method_line(doc, m)

    add_h(doc, "## backend/src/db.js", 2)
    for m in [
        "# getDb — Returns the singleton better-sqlite3 Database after ensuring the data directory exists and creating core tables if missing.",
        "# resetAllStoredData — Deletes all rows in core tables and clears backend/data/raw_html so each run can start clean.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## backend/src/httpError.js", 2)
    for m in [
        "# HttpError constructor — Builds an Error subclass carrying HTTP status and optional detail for consistent API error responses.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## backend/src/asyncRoute.js", 2)
    for m in [
        "# asyncRoute(handler) — Wraps an async Express handler so rejected promises call next(err) instead of crashing the connection.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## backend/src/scraper.js", 2)
    for m in [
        "# ensureRawDir — Creates backend/data/raw_html if it does not exist so HTML files can be written.",
        "# normalizeUrl — Trims input, adds https:// when no scheme is present, and rejects empty strings.",
        "# scrapeUrl — HTTP GETs a page, saves full HTML to disk, parses title/description/body text with cheerio, and inserts a scraped_pages row.",
        "# scrapeMany — Sequentially scrapes a list of URLs, collecting per-URL success objects or error messages without stopping the batch.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## backend/src/openaiService.js", 2)
    for m in [
        "# getClient — Lazily constructs a single OpenAI client or throws HttpError if OPENAI_API_KEY is missing.",
        "# jsonChat — Sends a system+user chat with JSON response_format (and optional max_tokens) and parses the assistant message into an object.",
        "# normalizeDeepDives — Coerces competitor_deep_dives into an array sized to the competitor list so multi-competitor comparison renders reliably.",
        "# buildCombinedContext — Calls the LLM to merge founder text with a scrape summary into one markdown brief plus assumptions and tags.",
        "# comparisonAndBrand — Produces exhaustive multi-competitor comparison tables plus one deep-dive card per competitor, along with brand strategy outputs.",
        "# roastStartup — Returns structured skeptical VC-style feedback including verdict, fixes, and positioning tweaks.",
        "# jobsFromCompetitors — Produces hiring thesis, organizational gaps, detailed job opportunities, and interview scorecard informed by competitors.",
        "# potentialVCsWithMatch — Suggests overlooked VC/investor targets with match_score 0–100 and outreach hooks.",
        "# strategicGapsAnalysis — Identifies product/team/market/fundraising gaps with severity and a 90-day closure plan.",
        "# startupIdeasTailored — Generates competitor-whitespace startup ideas with fit scores and validation steps.",
        "# collaborationOpportunities — Suggests ethical partnership types, pitches, and boundaries.",
        "# uniqueExtras — Outputs moat experiments, pivot seeds, community and pricing hypotheses, and metric ideas.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## backend/src/routes.js", 2)
    for m in [
        "# rowProfile — Reads one profiles row by primary key.",
        "# competitorRows — Lists all competitors for a profile ordered by creation time.",
        "# saveSnapshot — Inserts an analysis_snapshots row with JSON-serialized AI output.",
        "# loadContextBundle — Loads profile plus compact competitor summaries for prompts.",
        "# POST /api/profile — Upserts a profile, optionally scrapes startup URL, runs buildCombinedContext, updates combined_summary.",
        "# POST /api/competitors — Validates profile, scrapes competitor URLs, inserts competitor rows linked to scrape metadata.",
        "# GET /api/profile/:id — Returns profile row and competitor list or 404.",
        "# POST /api/analyze/:key — Dispatches comparison|vcs|jobs|gaps|ideas|roast|collaborations|extras to OpenAI and stores a snapshot.",
        "# GET /api/snapshots/:profileId — Returns recent snapshots with safe JSON parsing per row.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## frontend/src/api.js", 2)
    for m in [
        "# req — Performs fetch, parses JSON or fallback text errors, and throws on non-OK HTTP status.",
        "# saveProfile — POSTs startup text/URL (and optional profileId) to /api/profile.",
        "# addCompetitors — POSTs profileId and urls[] to /api/competitors.",
        "# getProfile — GETs /api/profile/:id.",
        "# runAnalysis — POSTs profileId to /api/analyze/:key.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## frontend/src/components/AnalysisResult.jsx", 2)
    for m in [
        "# ProseMarkdown — Renders markdown strings (including GFM tables) into readable typography.",
        "# ChipList — Renders a list of short strings as UI chips.",
        "# SeverityBadge — Renders low/medium/high severity as a colored badge.",
        "# MatchScoreBadge — Renders 0–100 match/fit scores as a colored badge.",
        "# FallbackTree — Renders unknown JSON shapes into nested key/value blocks without showing raw JSON first.",
        "# asDeepDiveList — Normalizes competitor_deep_dives into an array so the UI can render one card per competitor.",
        "# ComparisonView — Renders deep multi-competitor comparison tables and per-competitor deep-dive cards.",
        "# RoastView — Renders VC-style roast output with severity and a one-liner upgrade pullquote.",
        "# JobsView — Renders organizational gaps, roles, detailed job opportunities, and interview scorecard.",
        "# VcView — Renders ranked missed VCs with match scores, hooks, and diligence prep.",
        "# GapsView — Renders gap categories with severity plus a 90-day plan markdown block.",
        "# IdeasView — Renders tailored startup ideas with fit scores, risks, and validation steps.",
        "# AnalysisResult — Switches to the right view by panelKey and keeps raw JSON in a collapsible details block.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## frontend/src/App.jsx", 2)
    for m in [
        "# App (default export) — Holds form state, profileId in localStorage, calls API helpers, and renders analysis panels.",
        "# EMPTY_PANELS — Central reset template used to clear analysis outputs on new runs or profile changes.",
        "# persistProfile — Writes profileId to React state and localStorage.",
        "# useEffect (startup) — Validates saved profileId on load and clears local state if the backend data was reset.",
        "# handleSaveProfile — Calls saveProfile and stores the API response as lastProfile.",
        "# handleAddCompetitors — Calls addCompetitors then refreshes getProfile for display.",
        "# run — Clears old outputs, calls runAnalysis for a given key, and shows only the current result to avoid stale mixed panels.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## frontend/src/main.jsx", 2)
    for m in [
        "# (bootstrap) — Mounts the App root component into #root under React StrictMode.",
    ]:
        add_method_line(doc, m)

    add_h(doc, "## frontend/vite.config.js", 2)
    for m in [
        "# defineConfig — Enables React plugin and proxies /api to 127.0.0.1:3001 with extended timeouts and proxy error JSON fallback.",
    ]:
        add_method_line(doc, m)

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
